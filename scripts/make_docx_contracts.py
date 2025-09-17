#!/usr/bin/env python3
"""
Convert the provided PDF template to DOCX and add branded header/footer.
Outputs two editable files:
 - contracts/monthly-bookkeeping-agreement.docx
 - contracts/catch-up-bookkeeping-agreement.docx

Dependencies: pdf2docx, python-docx, cairosvg
Install: python3 -m pip install --user pdf2docx python-docx cairosvg
Run: python3 scripts/make_docx_contracts.py
"""
import os
import sys
from pathlib import Path

TEMPLATE_PDF = Path('lbooks-contract_template.pdf')
LOGO_SVG = Path('assets/img/favicon.svg')
LOGO_PNG = Path('assets/img/favicon-128.png')
BASE_DOCX = Path('contracts/_base-from-pdf.docx')
OUT1 = Path('contracts/monthly-bookkeeping-agreement.docx')
OUT2 = Path('contracts/catch-up-bookkeeping-agreement.docx')

def svg_to_png(svg_path: Path, png_path: Path, size: int = 128):
    """Render SVG to PNG without system cairo by routing through PDF.
    1) Draw SVG into a 1-page PDF using svglib+reportlab.
    2) Rasterize that PDF page to PNG using PyMuPDF.
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.graphics import renderPDF
    from svglib.svglib import svg2rlg
    import fitz  # PyMuPDF
    import io

    # Load and scale drawing
    drawing = svg2rlg(str(svg_path))
    scale = size / max(getattr(drawing, 'height', 24) or 24, 1)

    # Render to in-memory PDF
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.saveState()
    c.translate(36, 36)
    c.scale(scale, scale)
    renderPDF.draw(drawing, c, 0, 0)
    c.restoreState()
    c.showPage(); c.save()

    # Rasterize PDF page to PNG
    buf.seek(0)
    doc = fitz.open(stream=buf.getvalue(), filetype='pdf')
    page = doc.load_page(0)
    pix = page.get_pixmap(alpha=False)
    pix.save(str(png_path))

def convert_pdf_to_docx(pdf_path: Path, docx_path: Path):
    from pdf2docx import Converter
    c = Converter(str(pdf_path))
    try:
        c.convert(str(docx_path), start=0, end=None)
    finally:
        c.close()

def brand_docx(docx_in: Path, docx_out: Path, logo_png: Path):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document(str(docx_in))
    for section in doc.sections:
        header = section.header
        footer = section.footer

        # Clear existing header content
        header.is_linked_to_previous = False
        for p in list(header.paragraphs):
            p.clear()

        # Header: logo + brand + contact
        table = header.add_table(rows=1, cols=2, width=section.page_width)
        table.autofit = True
        left = table.cell(0, 0)
        right = table.cell(0, 1)

        # Left cell: logo + brand name
        p = left.paragraphs[0]
        run = p.add_run()
        try:
            run.add_picture(str(logo_png), width=Inches(0.33))
        except Exception:
            pass
        run = p.add_run('  Logical Books')
        run.font.bold = True
        run.font.size = Pt(11)

        # Right cell: contact line, right aligned
        pr = right.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = pr.add_run('logicalbooks.io  •  (336) 858‑3549  •  info@logicalbooks.com')
        rr.font.size = Pt(9)

        # Footer: Page X aligned right
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p.clear()
        pf = footer.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = pf.add_run('Page ')
        # Insert PAGE field
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
        pf._p.append(fld)

    doc.save(str(docx_out))

def main():
    if not TEMPLATE_PDF.exists():
        print(f"Missing {TEMPLATE_PDF}", file=sys.stderr)
        sys.exit(1)

    OUT1.parent.mkdir(parents=True, exist_ok=True)

    # Ensure raster logo
    if LOGO_SVG.exists() and not LOGO_PNG.exists():
        svg_to_png(LOGO_SVG, LOGO_PNG, size=128)

    # Convert base once
    if not BASE_DOCX.exists():
        convert_pdf_to_docx(TEMPLATE_PDF, BASE_DOCX)

    # Brand two copies
    brand_docx(BASE_DOCX, OUT1, LOGO_PNG if LOGO_PNG.exists() else LOGO_SVG)
    brand_docx(BASE_DOCX, OUT2, LOGO_PNG if LOGO_PNG.exists() else LOGO_SVG)

    print(f"Wrote {OUT1}")
    print(f"Wrote {OUT2}")

if __name__ == '__main__':
    main()
